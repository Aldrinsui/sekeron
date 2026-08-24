#!/usr/bin/env python3
"""
Sekeron Stage 3 - Artist Intelligence: Profile Text Extraction
===================================================================

Deterministic, local, non-LLM extraction of plain text from each artist's
profile document (.docx or .txt), read from the Stage 1 dataset manifest.

Produces:
    generated/artist_profiles.json

IMPORTANT - CLAIMS, NOT CAPABILITIES:
Every profile_text extracted here is CLAIMED information written by or
about the artist. It has NOT been verified against any media evidence.
Downstream stages MUST NOT treat any statement in profile_text as a
demonstrated capability. It is input to the "profile claims" side of the
evidence model only - see decision_note.md / README.md for the full
claims-vs-demonstrated distinction. This script does not interpret,
score, or classify the text in any way; it only extracts it verbatim.

READ-ONLY GUARANTEE (same as Stage 1 / Stage 2):
This script only opens source files for reading (python-docx's Document()
call and Python's built-in open() in read mode). It never writes, renames,
or modifies anything under --dataset-root. --output-dir is checked at
startup and the script refuses to run if it resolves inside --dataset-root.

NO LLM calls. NO capability inference. NO ranking. NO embeddings.

Usage:
    python3 scripts/extract_profiles.py \\
        --dataset-root "/path/to/Data set" \\
        --manifest generated/dataset_manifest.json \\
        --output-dir generated

Requires: Python 3.9+, python-docx (see requirements.txt).
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install -r requirements.txt", file=sys.stderr)
    sys.exit(1)

EXPECTED_ARTIST_COUNT = 15


def extract_docx_text(path: Path) -> str:
    """Extract plain text from a .docx file: paragraphs and table cell text,
    in document order, joined with newlines. Read-only - python-docx's
    Document() does not modify the source file."""
    doc = Document(str(path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return "\n".join(parts)


def extract_txt_text(path: Path) -> str:
    """Extract plain text from a .txt file. Tries UTF-8 first, falls back to
    latin-1 (never raises on encoding alone) so a single odd file doesn't
    abort the whole run - any real read failure is still caught upstream."""
    try:
        return path.read_text(encoding="utf-8").strip()
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1").strip()


def extract_profile(dataset_root: Path, profile_record: dict) -> dict:
    rel_path = profile_record["relative_path"]
    full_path = dataset_root / rel_path
    ext = full_path.suffix.lower()

    result = {
        "source_relative_path": rel_path,
        "profile_text": None,
        "extraction_status": "ok",
        "extraction_error": None,
    }

    if not full_path.exists():
        result["extraction_status"] = "file_not_found"
        result["extraction_error"] = f"File listed in manifest but not found on disk: {rel_path}"
        return result

    try:
        if ext == ".docx":
            text = extract_docx_text(full_path)
        elif ext == ".txt":
            text = extract_txt_text(full_path)
        else:
            result["extraction_status"] = "unsupported_format"
            result["extraction_error"] = f"Unsupported profile document extension: {ext}"
            return result

        if not text:
            result["extraction_status"] = "empty_text"
            result["extraction_error"] = "Document opened successfully but no extractable text was found."
        result["profile_text"] = text if text else None

    except Exception as e:
        result["extraction_status"] = "extraction_failed"
        result["extraction_error"] = f"{type(e).__name__}: {e}"

    return result


def main():
    parser = argparse.ArgumentParser(description="Extract plain text from artist profile documents (read-only).")
    parser.add_argument("--dataset-root", required=True, type=Path)
    parser.add_argument("--manifest", required=True, type=Path,
                         help="Path to Stage 1 dataset_manifest.json")
    parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args()

    dataset_root = args.dataset_root.resolve()
    output_dir = args.output_dir.resolve()

    if not dataset_root.exists():
        print(f"ERROR: dataset root does not exist: {dataset_root}", file=sys.stderr)
        sys.exit(1)
    if not args.manifest.exists():
        print(f"ERROR: manifest not found: {args.manifest}", file=sys.stderr)
        sys.exit(1)

    try:
        output_dir.relative_to(dataset_root)
        print(f"ERROR: --output-dir ({output_dir}) is inside --dataset-root ({dataset_root}). Refusing to run.",
              file=sys.stderr)
        sys.exit(1)
    except ValueError:
        pass

    with open(args.manifest, "r", encoding="utf-8") as f:
        dataset_manifest = json.load(f)

    output_dir.mkdir(parents=True, exist_ok=True)

    records = []
    warnings = []

    for artist in dataset_manifest["artists"]:
        artist_id = artist["artist_id"]
        category = artist["category"]
        profile_files = artist.get("profile_files", [])

        if not profile_files:
            records.append({
                "artist_id": artist_id,
                "category": category,
                "source_relative_path": None,
                "profile_text": None,
                "extraction_status": "no_profile_document",
                "extraction_error": "No profile_files entry found for this artist in the Stage 1 manifest.",
            })
            warnings.append(f"{artist_id}: no profile document found.")
            continue

        if len(profile_files) > 1:
            warnings.append(
                f"{artist_id}: {len(profile_files)} profile documents found; extracting all of them "
                f"as separate records (Stage 1 manifest did not indicate a single canonical profile file)."
            )

        for pf in profile_files:
            extraction = extract_profile(dataset_root, pf)
            record = {
                "artist_id": artist_id,
                "category": category,
                **extraction,
            }
            records.append(record)
            if extraction["extraction_status"] != "ok":
                warnings.append(f"{artist_id}: {extraction['extraction_status']} - {extraction['extraction_error']}")

    # --- Validation: all 15 expected artists represented, each with at least
    # one profile record whose extraction_status is 'ok'. ---
    artist_ids_in_manifest = {a["artist_id"] for a in dataset_manifest["artists"]}
    artist_ids_with_ok_extraction = {
        r["artist_id"] for r in records if r["extraction_status"] == "ok"
    }
    artist_ids_with_any_record = {r["artist_id"] for r in records}

    validation = {
        "expected_artist_count": EXPECTED_ARTIST_COUNT,
        "artist_count_in_manifest": len(artist_ids_in_manifest),
        "artists_with_any_profile_record": sorted(artist_ids_with_any_record),
        "artists_with_successful_extraction": sorted(artist_ids_with_ok_extraction),
        "artists_missing_from_manifest": [],
        "artists_with_failed_or_missing_extraction": sorted(artist_ids_in_manifest - artist_ids_with_ok_extraction),
        "all_15_represented": False,
        "all_15_successfully_extracted": False,
    }

    if len(artist_ids_in_manifest) != EXPECTED_ARTIST_COUNT:
        validation["artists_missing_from_manifest"] = (
            f"Manifest contains {len(artist_ids_in_manifest)} artists, expected {EXPECTED_ARTIST_COUNT}."
        )
        warnings.append(validation["artists_missing_from_manifest"])

    validation["all_15_represented"] = (
        len(artist_ids_with_any_record) == EXPECTED_ARTIST_COUNT
        and artist_ids_with_any_record == artist_ids_in_manifest
    )
    validation["all_15_successfully_extracted"] = (
        len(artist_ids_with_ok_extraction) == EXPECTED_ARTIST_COUNT
    )

    if not validation["all_15_represented"]:
        warnings.append(
            f"NOT all {EXPECTED_ARTIST_COUNT} expected artists have a profile record. "
            f"Present: {sorted(artist_ids_with_any_record)}"
        )
    if not validation["all_15_successfully_extracted"]:
        warnings.append(
            f"NOT all {EXPECTED_ARTIST_COUNT} artists have a successfully extracted profile. "
            f"Failed/missing: {validation['artists_with_failed_or_missing_extraction']}"
        )

    output = {
        "generated_by": "scripts/extract_profiles.py",
        "manual_edits": False,
        "dataset_root": str(dataset_root),
        "note": "profile_text is CLAIMED information only. It must never be treated as a demonstrated "
                "capability by downstream stages without corroborating media evidence.",
        "validation": validation,
        "warnings": warnings,
        "artist_profiles": records,
    }

    out_path = output_dir / "artist_profiles.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)

    print(f"Wrote: {out_path}")
    print(f"Artists in manifest: {len(artist_ids_in_manifest)}")
    print(f"Profile records extracted: {len(records)}")
    print(f"All {EXPECTED_ARTIST_COUNT} artists represented: {validation['all_15_represented']}")
    print(f"All {EXPECTED_ARTIST_COUNT} artists successfully extracted: {validation['all_15_successfully_extracted']}")
    if warnings:
        print(f"\n{len(warnings)} warning(s):")
        for w in warnings:
            print(f"  - {w}")


if __name__ == "__main__":
    main()
